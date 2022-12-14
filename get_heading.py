"""
def iter_headings(paragraphs):
    for paragraph in paragraphs:
        if paragraph.style.name.startswith('Heading'):
            yield paragraph
for heading in inter_headings(document.paragraphs):
    print heading.text
"""
import docx 
import aspose
# documents = docx.Document('sample.docx')

# def iter_headings(test_word):
#     for texts in test_word:
#         if texts.style.name.startswith('Heading 1'):
#             yield texts

# for heading in iter_headings(documents.test_word):
#     print(heading)

# doc=docx.Document('sample.docx')


# def iter_heading(paragraphs):
#     for paragraph in paragraphs:
#         if paragraph.style.name.startswith('Heading 1'):
#             yield paragraph

# for heading in iter_heading(doc.paragraphs):
#     print("Text of word",heading)

# import win32com.client as win32
# word = win32.Dispatch("Word.Application")
# word.Visible = 0
# word.Documents.Open("F:\Learn-And-Practice\sample.docx")
# doc = word.ActiveDocument

# for word in doc.Words:
#     print(word)
# Load document
import aspose.words as aw
import helps
from helps import paragraphs_by_style_name
import helpers
from helpers import extract_content
from helps import verify_parameter_nodes
from helps import generate_document
# Load document
doc = aw.Document("sample.docx")

# # Gather a list of the paragraphs using the respective heading styles.
# parasStyleHeading1 = paragraphs_by_style_name(doc, "Heading 1")
# parasStyleHeading3 = paragraphs_by_style_name(doc, "Heading 3")

# # Use the first instance of the paragraphs with those styles.
# startPara1 = parasStyleHeading1[0]
# endPara1 = parasStyleHeading3[0]

# # Extract the content between these nodes in the document. Don't include these markers in the extraction.
# extractedNodes = extract_content(startPara1, endPara1, False)

# # Generate document containing extracted content.
# dstDoc = generate_document(doc, extractedNodes)

# # Save document.
# dstDoc.save("extract_content_between_paragraphs_based_on-Styles.docx")

# Define starting and ending paragraphs.
startPara = doc.first_section.body.get_child(aw.NodeType.PARAGRAPH, 6, True).as_paragraph()
endPara = doc.first_section.body.get_child(aw.NodeType.PARAGRAPH, 10, True).as_paragraph()

# Extract the content between these paragraphs in the document. Include these markers in the extraction.
extractedNodes = extract_content(startPara, endPara, True)

# Generate document containing extracted content.
dstDoc = generate_document(doc, extractedNodes)

# Save document.
dstDoc.save("extract_content.docx")